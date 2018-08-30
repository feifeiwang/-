# coding=gbk
import docx
# from docx import Document
# from docx.shared import Inches
import re
import os
import datetime

print('Please check the folder  D:\\baoxiao')
print('Please check the file  D:\\����.docx')
change = input('Press any key to continue..')


def readDocx(docName):
    doc = docx.Document(docName)
    docto = docx.Document('D:\\����.docx')
    table1 = doc.tables
    table2 = docto.tables
    tab = table1[0]
    huizong = table2[0]
    hang = (len(tab.rows))
    lie = (len(tab.columns))
    ix = 2
    zong = 0
    print(tab.rows[2].cells[0].text)
    zuihou = tab.rows[hang - 1].cells[0].text
    jiegou = re.search(r'\d+.\d*', zuihou, re.M | re.I)
    jisuan = float(jiegou.group(0))
    xm = 0
    while ix < hang - 1:
        row = tab.rows[ix]
        try:
            shu = float(row.cells[2].text)
        except:
            pass
        else:
            zong = zong + shu
        name = row.cells[3].text
        value = row.cells[2].text
        if name:
            xm = xm + 1
            print(name, '     ', value)
            huizong.add_row()
            hang2 = (len(huizong.rows))
            huirow = huizong.rows[hang2 - 1]
            lieshu = 1
            if ix == 2:
                huirow.cells[0].text = row.cells[0].text

            while lieshu < 7:
                huirow.cells[lieshu].text = row.cells[lieshu].text
                lieshu = lieshu + 1
        ix = ix + 1
    huizong.add_row()
    huirow = huizong.rows[hang2]
    huirow.cells[0].text = zuihou
    # �ϲ���Ԫ��
    aa = huirow.cells[0]
    iy = 1
    while iy < 7:
        bb = huirow.cells[iy]
        iy = iy + 1
        aa.merge(bb)
    huicol = huizong.columns[0]
    cc = huicol.cells[hang2 - (xm)]
    if xm > 1:
        iz = 1
        while iz < (xm):
            dd = huicol.cells[hang2 - (xm) + iz]
            cc.merge(dd)
            iz = iz + 1

    wucha = abs(jisuan - zong)
    if wucha < 0.001:
        abc = '            True'
    else:
        abc = '           !!!!!    ��������!!!!!!!'
        print('�ܶԪ����', jisuan, ' ����ֵ��', zong, abc)
        return -1
    print('�ܶԪ����', jisuan, ' ����ֵ��', zong, abc)

    docto.save('D:\\����.docx')
    return zong


pathDir = os.listdir('D:\\baoxiao')
x = 0
zonge = 0

for allDir in pathDir:
    x = x + 1
    child = os.path.join('%s%s' % ('D:\\\\baoxiao\\\\', allDir))
    print(child[13:])
    if child[-4:]!='docx':
        print('-----------------------------------------------------')
        print('----��֧��.doc��ȡ�����޸��ļ���ʽ����������������---')
        print('-----------------------------------------------------')
        break
        break
print('��ȷ�����ļ�����', x)
name = input('Press any key to continue..')

x = 0

for allDir in pathDir:
    x = x + 1
    print('�ļ���', x)
    child = os.path.join('%s%s' % ('D:\\\\baoxiao\\\\', allDir))
    print(child[13:])
    try:
        dan = readDocx(child)
    except:

        print('-----------------------------------------------------')
        print('----�����������������ĵ��д������������������---')
        print('-----------------------------------------------------')
        break
    if dan == -1:
        print('!!!!!!�����ļ�!!!!!!')
        print(child)
        print('!!!!!!�����ļ�!!!!!!')
        break
    zonge = zonge + dan
docto = docx.Document('D:\\����.docx')
zonge = round(zonge, 2)
table3 = docto.tables
huizong2 = table3[0]
huizong2.add_row()
hang3 = (len(huizong2.rows))
huirow2 = huizong2.rows[hang3 - 1]
now = datetime.datetime.now()
now.strftime('%Y-%m-%d %H:%M:%S')
zongjie = '�ϼƣ�' + str(zonge) + '   �������� ���ϱ�����ĿĿ����ȷ��������ʵ      ʱ�䣺' + now.strftime(
    '%Y-%m-%d %H:%M:%S') + '  \n�����Ա��        ���⸺���ˣ���׼����             �����ˣ�        '
print('####################################################################################')
print(zongjie)
huirow2.cells[0].text = zongjie
aa = huirow2.cells[0]
iy = 1
while iy < 7:
    bb = huirow2.cells[iy]
    iy = iy + 1
    aa.merge(bb)
docto.save('D:\\����.docx')

change = input('Press any key to qiut!!')